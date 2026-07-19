from __future__ import annotations

import argparse
import hashlib
import json
import re
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


PLACEHOLDER = re.compile(
    r"(?:\.{4,}|_{4,}|\[[^\]]{1,100}\]|\b(?:cumplimentar|rellenar|a completar|pendiente)\b)",
    re.IGNORECASE,
)
TEMPLATE_NAME = re.compile(r"\b(?:plantilla|modelo|formulario)\b", re.IGNORECASE)
COMPLETED_PATH = re.compile(
    r"(?:firmad|signed|registrad|entregad|justificante|documentos para entregar)", re.IGNORECASE
)
PERSONAL = re.compile(
    r"\b(?:DNI|NIE|nombre y apellidos|persona de contacto|representante|correo|tel[eé]fono|firma)\b",
    re.IGNORECASE,
)
SENSITIVE = re.compile(
    r"\b(?:persona becada|persona beneficiaria|menor(?:es)?|salud|discapacidad|migrante|violencia)\b",
    re.IGNORECASE,
)
SUPPORTED = {".docx", ".pdf", ".xlsx"}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Inventaría plantillas privadas sin copiar ni publicar su contenido."
    )
    parser.add_argument("--corpus", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path)
    parser.add_argument("--tenant", required=True)
    return parser.parse_args()


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for block in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def read_docx(path: Path) -> tuple[str, dict]:
    document = Document(str(path))
    paragraphs = [item.text for item in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    empty_cells = sum(not value.strip() for value in cells)
    return "\n".join(paragraphs + cells), {
        "format": "docx",
        "tables": len(document.tables),
        "cells": len(cells),
        "empty_cell_ratio": round(empty_cells / max(1, len(cells)), 3),
        "form_fields": 0,
    }


def read_pdf(path: Path) -> tuple[str, dict]:
    reader = PdfReader(str(path), strict=False)
    text = "\n".join((page.extract_text() or "") for page in reader.pages[:20])
    fields = reader.get_fields() or {}
    empty_fields = sum(not (field or {}).get("/V") for field in fields.values())
    return text, {
        "format": "pdf",
        "pages_sampled": min(20, len(reader.pages)),
        "form_fields": len(fields),
        "empty_form_fields": empty_fields,
        "cells": 0,
        "empty_cell_ratio": 0,
    }


def read_xlsx(path: Path) -> tuple[str, dict]:
    book = load_workbook(path, read_only=True, data_only=True)
    values: list[str] = []
    cells = empty = 0
    for sheet in book.worksheets[:12]:
        max_row = min(250, sheet.max_row or 1)
        max_col = min(30, sheet.max_column or 1)
        for row in sheet.iter_rows(max_row=max_row, max_col=max_col, values_only=True):
            for value in row:
                cells += 1
                if value in (None, ""):
                    empty += 1
                else:
                    values.append(str(value))
    return "\n".join(values), {
        "format": "xlsx",
        "sheets_sampled": min(12, len(book.worksheets)),
        "form_fields": 0,
        "cells": cells,
        "empty_cell_ratio": round(empty / max(1, cells), 3),
    }


def document_kind(path: Path, text: str) -> str:
    haystack = f"{path.name}\n{text[:20_000]}".lower()
    if "propuesta económica" in haystack or "presupuesto" in path.name.lower():
        return "economic_offer"
    if "carta de aval" in haystack:
        return "endorsement_letter"
    if "persona becada" in haystack or "informe seguimiento" in haystack:
        return "participant_report"
    if "plantilla evidencias personal" in haystack or "plantilla personal" in haystack:
        return "staff_evidence"
    if "acuerdo" in haystack or "convenio" in haystack:
        return "agreement"
    if re.search(r"\b(?:solicitud|memoria|anexo)\b", haystack):
        return "application_or_memory"
    return "unmapped"


def field_policy(text: str) -> tuple[list[str], list[str]]:
    safe: list[str] = []
    blocked: list[str] = []
    rules = [
        (r"\b(?:entidad|empresa|fundaci[oó]n)\b", "legal_name", safe),
        (r"\b(?:CIF|NIF)\b", "tax_id", safe),
        (r"\bdomicilio\b", "registered_address", safe),
        (r"\brepresentante\b", "representative_name", blocked),
        (r"\b(?:DNI|NIE)\b", "personal_id", blocked),
        (r"\b(?:importe|euros|presupuesto)\b", "amounts", blocked),
        (r"\bfecha\b", "dates", blocked),
        (r"\bfirma\b", "signature", blocked),
        (r"\b(?:persona becada|persona beneficiaria|datos de la persona|nombre y apellidos)\b", "participant_data", blocked),
    ]
    for pattern, key, target in rules:
        if re.search(pattern, text, re.IGNORECASE):
            target.append(key)
    return sorted(set(safe)), sorted(set(blocked))


def classify(path: Path, text: str, structure: dict) -> dict:
    placeholders = len(PLACEHOLDER.findall(text))
    name_hint = bool(TEMPLATE_NAME.search(path.name))
    empty_form = structure["cells"] >= 8 and structure["empty_cell_ratio"] >= 0.25
    if structure["format"] == "pdf":
        candidate = bool(structure.get("empty_form_fields") or (name_hint and placeholders >= 1))
    elif structure["format"] == "xlsx":
        candidate = bool(name_hint or placeholders >= 2)
    else:
        candidate = bool(placeholders >= 2 or name_hint or (empty_form and len(text) < 8_000))
        if len(text) > 30_000 and placeholders < 5 and structure["empty_cell_ratio"] < 0.1:
            candidate = False
    completed_copy = bool(COMPLETED_PATH.search(str(path)))
    personal = bool(PERSONAL.search(text))
    sensitive = bool(SENSITIVE.search(text))
    safe_fields, blocked_fields = field_policy(text)

    kind = document_kind(path, text)
    if not candidate or completed_copy:
        decision = "reference_only"
    elif sensitive or kind == "participant_report":
        decision = "blocked_sensitive"
    elif kind in {"staff_evidence", "endorsement_letter"} or (kind == "agreement" and personal):
        decision = "manual_only"
    elif safe_fields:
        decision = "map_before_prefill"
    elif personal:
        decision = "manual_only"
    else:
        decision = "manual_review"

    data_class = "sensitive" if sensitive else "personal" if personal else "internal"
    return {
        "candidate": candidate and not completed_copy,
        "kind": kind,
        "decision": decision,
        "data_class": data_class,
        "placeholder_count": placeholders,
        "safe_field_keys": safe_fields,
        "blocked_field_keys": blocked_fields,
        "structure": structure,
    }


def refine_candidates(documents: list[dict]) -> None:
    seen_hashes: dict[str, str] = {}
    for item in documents:
        if not item["candidate"]:
            continue
        fields = int(item["structure"].get("form_fields") or 0)
        empty = int(item["structure"].get("empty_form_fields") or 0)
        if fields and empty / fields < 0.8:
            item["candidate"] = False
            item["decision"] = "reference_only_filled"
            item["reference_reason"] = "form_already_filled"
            continue
        source_hash = item["source_sha256"]
        if source_hash in seen_hashes:
            item["candidate"] = False
            item["decision"] = "duplicate_reference"
            item["duplicate_of_document_id"] = seen_hashes[source_hash]
            continue
        seen_hashes[source_hash] = item["document_id"]


def main() -> None:
    config = parse_args()
    corpus = config.corpus.resolve()
    output = config.out.resolve()
    if not corpus.is_dir():
        raise ValueError("El corpus autorizado no existe o no es una carpeta.")
    if output.is_relative_to(corpus):
        raise ValueError("El inventario no puede escribirse dentro del corpus fuente.")

    readers = {".docx": read_docx, ".pdf": read_pdf, ".xlsx": read_xlsx}
    documents: list[dict] = []
    errors: list[dict] = []
    for path in sorted(corpus.rglob("*")):
        suffix = path.suffix.lower()
        if not path.is_file() or suffix not in SUPPORTED:
            continue
        relative = str(path.relative_to(corpus))
        try:
            text, structure = readers[suffix](path)
            documents.append({
                "document_id": sha256(path)[:16],
                "source_sha256": sha256(path),
                "relative_path": relative,
                "extension": suffix,
                **classify(path, text, structure),
            })
        except Exception as error:
            errors.append({"relative_path": relative, "extension": suffix, "error": type(error).__name__})

    refine_candidates(documents)
    decisions = Counter(item["decision"] for item in documents)
    report = {
        "tenant": config.tenant,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "privacy": {
            "processing": "local_only",
            "external_ai_calls": 0,
            "content_copied_to_report": False,
            "cross_tenant_reuse": False,
            "human_review_required": True,
        },
        "metrics": {
            "documents_scanned": len(documents),
            "template_candidates": sum(item["candidate"] for item in documents),
            "decisions": dict(sorted(decisions.items())),
            "extraction_errors": len(errors),
        },
        "documents": documents,
        "errors": errors,
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"ok": True, "output": str(output), **report["metrics"]}, ensure_ascii=False))


if __name__ == "__main__":
    main()
