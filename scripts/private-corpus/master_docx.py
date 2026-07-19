from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_font(run, size=11, bold=False, color="24233A"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def fix_table_geometry(table, widths: list[int]):
    total = sum(widths)
    properties = table._tbl.tblPr
    for tag, value in [("w:tblW", total), ("w:tblInd", 120)]:
        node = properties.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            properties.append(node)
        node.set(qn("w:type"), "dxa")
        node.set(qn("w:w"), str(value))
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_properties = cell._tc.get_or_add_tcPr()
            tc_w = tc_properties.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_properties.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def build_docx(path: Path, tenant: str, facts: dict, selected: dict, metrics: dict, fields: dict):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after in [("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 12, 8, 4)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("414F8E")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("INSERTIA · CONOCIMIENTO PRIVADO DEL TENANT"), 8, True, "566964")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Documento de trabajo · No presentar sin revisión humana"), 8, False, "737889")

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("PROPUESTA DOCUMENTAL TRAZABLE"), 10, True, "3D866A")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    set_font(title.add_run("Plantilla maestra propuesta"), 24, True, "24233A")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    set_font(subtitle.add_run(tenant), 13, False, "56566C")

    callout = doc.add_table(rows=1, cols=1)
    callout.autofit = False
    callout.columns[0].width = Inches(6.5)
    shade(callout.cell(0, 0), "FFF6DF")
    callout.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = callout.cell(0, 0).paragraphs[0]
    set_font(paragraph.add_run("PENDIENTE DE APROBACIÓN · "), 10, True, "8A5A12")
    set_font(paragraph.add_run("Contenido recuperado localmente. No contiene datos personales detectados y conserva evidencia por campo."), 10, False, "5B6073")
    fix_table_geometry(callout, [9360])

    doc.add_heading("1. Datos institucionales", level=1)
    for fact in facts.values():
        paragraph = doc.add_paragraph()
        set_font(paragraph.add_run(f"{fact['field']}: "), 11, True)
        set_font(paragraph.add_run(fact["value"] or "[PENDIENTE DE APORTAR]"), 11)
        if fact["conflicts"]:
            note = doc.add_paragraph()
            note.paragraph_format.left_indent = Inches(.2)
            set_font(note.add_run("Conflicto detectado: existen variantes y debe confirmarse el valor antes de aprobar."), 9, False, "9B650E")

    doc.add_heading("2. Contenido reutilizable propuesto", level=1)
    for key, (label, _) in fields.items():
        doc.add_heading(label, level=2)
        item = selected.get(key)
        if not item:
            set_font(doc.add_paragraph().add_run("[PENDIENTE: no se encontró evidencia segura suficiente en el corpus.]"), 10, True, "9B650E")
            continue
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = True
        set_font(paragraph.add_run(item["text"]), 11)
        source = doc.add_paragraph()
        source.paragraph_format.space_before = Pt(4)
        source.paragraph_format.space_after = Pt(4)
        source.paragraph_format.keep_together = True
        set_font(source.add_run(f"Evidencia: {item['source']} · documento {item['document_id']} · propuesta sin aprobar"), 8, False, "566964")

    doc.add_heading("3. Campos que nunca se completan sin control", level=1)
    for text in ["Representante y firma", "Importe solicitado y presupuesto vigente", "Fechas y plazo de ejecución", "Certificados y vigencia", "Datos personales de equipo o participantes"]:
        set_font(doc.add_paragraph(style="List Bullet").add_run(text), 11)

    doc.add_heading("4. Control de cobertura", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    for index, width in enumerate([Inches(2.7), Inches(1.25), Inches(2.55)]):
        table.columns[index].width = width
    for cell, text in zip(table.rows[0].cells, ["Bloque", "Estado", "Evidencia"]):
        shade(cell, "E4F3E9")
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_font(cell.paragraphs[0].add_run(text), 10, True, "285C3D")
    for key, (label, _) in fields.items():
        row = table.add_row().cells
        item = selected.get(key)
        values = [label, "Propuesto" if item else "Pendiente", item["document_id"] if item else "Sin evidencia"]
        for cell, text in zip(row, values):
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            set_font(cell.paragraphs[0].add_run(text), 9)
    fix_table_geometry(table, [3888, 1800, 3672])

    summary = doc.add_paragraph()
    summary.paragraph_format.space_before = Pt(10)
    set_font(summary.add_run(f"Cobertura automática inicial: {metrics['filled_sections']}/{metrics['total_sections']} bloques ({metrics['coverage_percent']}%). "), 10, True)
    set_font(summary.add_run("La cobertura mide propuestas, no aprobación ni aptitud para presentar."), 10)

    doc.core_properties.author = "Insertia local"
    doc.core_properties.last_modified_by = "Insertia local"
    doc.core_properties.comments = "Generado sin servicios externos; tenant privado; revisión humana obligatoria."
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
