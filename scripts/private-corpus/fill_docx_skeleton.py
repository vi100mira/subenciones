from __future__ import annotations

import argparse
import hashlib
import json
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ENTITY_BLOCK = re.compile(
    r"(empresa/entidad\s+)(?:[.\u2026]{3,}|_+|\[[^\]]+\])"
    r"(\s+con\s+CIF\s*)(?:[.\u2026]{3,}|_+|\[[^\]]+\])",
    re.IGNORECASE,
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Completa solo campos institucionales seguros de un DOCX privado."
    )
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--facts-json", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path)
    parser.add_argument("--audit", required=True, type=Path)
    parser.add_argument("--tenant", required=True)
    parser.add_argument("--allow-proposed-drafts", action="store_true")
    return parser.parse_args()


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def set_text(paragraph, text: str) -> None:
    style = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)
    style.text = text


def insert_notice(document: Document) -> None:
    anchor = document.paragraphs[0]._p
    notice = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "FFF4D6")
    properties.append(shading)
    notice.append(properties)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    bold = OxmlElement("w:b")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "8A5A12")
    run_properties.extend([bold, color])
    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    text.text = (
        "BORRADOR PARCIAL · Se han propuesto únicamente razón social y CIF. "
        "Representante, domicilio personal, importes, fecha, cargo y firma "
        "requieren cumplimentación y aprobación humana."
    )
    run.extend([run_properties, text])
    notice.append(run)
    anchor.addprevious(notice)

    spacer = OxmlElement("w:p")
    notice.addnext(spacer)


def add_footer(document: Document) -> None:
    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(
            "Insertia · borrador privado del tenant · no presentar sin revisión humana"
        )
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(86, 105, 100)


def main() -> None:
    config = parse_args()
    template = config.template.resolve()
    facts_path = config.facts_json.resolve()
    output = config.out.resolve()
    audit_path = config.audit.resolve()
    if template == output:
        raise ValueError("El borrador no puede sobrescribir la plantilla.")

    payload = json.loads(facts_path.read_text(encoding="utf-8"))
    payload = payload.get("data") if payload.get("ok") is True else payload
    if payload.get("tenant") != config.tenant:
        raise ValueError("Los hechos no pertenecen al tenant solicitado.")
    facts = payload["institutional_facts"]
    legal_name = facts["legal_name"]
    tax_id = facts["tax_id"]
    if not legal_name.get("value") or not tax_id.get("value"):
        raise ValueError("Faltan razón social o CIF propuestos en el perfil maestro.")

    facts_approved = all(item.get("status") == "approved" for item in [legal_name, tax_id])
    if not facts_approved and not config.allow_proposed_drafts:
        raise ValueError("Las propuestas solo pueden usarse con --allow-proposed-drafts.")

    document = Document(str(template))
    matches = 0
    for paragraph in document.paragraphs:
        original = paragraph.text
        replaced, count = ENTITY_BLOCK.subn(
            rf"\g<1>{legal_name['value']}\g<2> {tax_id['value']}", original
        )
        if count:
            set_text(paragraph, replaced)
            matches += count

    if matches != 1:
        raise ValueError(
            f"Se esperaba exactamente un bloque empresa/CIF; encontrados: {matches}."
        )

    insert_notice(document)
    add_footer(document)
    document.core_properties.author = "Insertia local"
    document.core_properties.last_modified_by = "Insertia local"
    document.core_properties.comments = (
        "Documento privado; autocompletado parcial; revisión humana obligatoria."
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)

    audit = {
        "tenant": config.tenant,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "processing": "local_only",
        "external_ai_calls": 0,
        "template_sha256": digest(template),
        "output_sha256": digest(output),
        "filled": [
            {
                "field": "legal_name",
                "status": "verified_from_approved_fact" if legal_name.get("status") == "approved" else "proposed_pending_human_review",
                "evidence": legal_name.get("evidence", []),
            },
            {
                "field": "tax_id",
                "status": "verified_from_approved_fact" if tax_id.get("status") == "approved" else "proposed_pending_human_review",
                "evidence": tax_id.get("evidence", []),
            },
        ],
        "personal_values_reused": False,
        "blocked": [
            "representative_name",
            "representative_address",
            "amounts",
            "date",
            "role",
            "signature",
        ],
        "human_review_required": True,
        "ready_to_present_percentage": 0,
        "external_submission_allowed": False,
    }
    audit_path.parent.mkdir(parents=True, exist_ok=True)
    audit_path.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(
        json.dumps(
            {"ok": True, "fields_filled": matches + 1, "output": str(output), "audit": str(audit_path)},
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
