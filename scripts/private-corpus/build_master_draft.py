from __future__ import annotations

import argparse
import hashlib
import json
import re
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from master_docx import build_docx


FIELDS = {
    "mission": ("Misión y fines", ["misión", "fines fundacionales", "objeto social", "acompañar a personas"]),
    "trajectory": ("Trayectoria y capacidad", ["trayectoria", "años de experiencia", "experiencia en", "desde el año", "proyectos ejecutados"]),
    "territory": ("Implantación territorial", ["ámbito territorial", "comunitat valenciana", "provincia de", "municipios", "implantación territorial"]),
    "collectives": ("Colectivos destinatarios", ["personas destinatarias", "colectivos", "riesgo de exclusión", "vulnerabilidad", "dificultades de acceso al empleo"]),
    "methodology": ("Metodología de intervención", ["metodología", "itinerario", "diagnóstico individualizado", "orientación", "intermediación laboral"]),
    "team": ("Equipo y capacidad operativa", ["equipo técnico", "recursos humanos", "personal técnico", "atención directa", "coordinación"]),
    "alliances": ("Alianzas y trabajo en red", ["entidades colaboradoras", "empresas colaboradoras", "trabajo en red", "servicios sociales", "acuerdos de colaboración"]),
    "evaluation": ("Evaluación e indicadores", ["indicadores", "evaluación", "resultados esperados", "seguimiento", "mejora continua"]),
}

EMAIL = re.compile(r"\b[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}\b")
PHONE = re.compile(r"(?<!\d)(?:\+34\s*)?(?:\d[\s.-]*){9}(?!\d)")
PERSON_ID = re.compile(r"\b(?:[XYZ]\d{7}[A-Z]|\d{8}[A-Z])\b", re.I)
IBAN = re.compile(r"\bES\s*\d{2}(?:\s*\d{4}){5}\b", re.I)
PERSON_CUES = re.compile(r"\b(?:D\.?/?Dª|Dña\.?|DNI|NIE|persona de contacto|teléfono|correo electrónico|firma(?:do)? por|nombre y apellidos)\b", re.I)
PLACEHOLDER = re.compile(r"\.{5,}|_{4,}|\[[^\]]{2,80}\]|\b(?:cumplimentar|rellenar|indicar|pendiente)\b", re.I)


def args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Construye una propuesta documental privada y trazable sin servicios externos.")
    parser.add_argument("--corpus", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path)
    parser.add_argument("--tenant", required=True)
    parser.add_argument("--entity-name", required=True)
    return parser.parse_args()


def extract(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        doc = Document(str(path)); blocks = [p.text for p in doc.paragraphs]
        blocks.extend(" | ".join(cell.text for cell in row.cells) for table in doc.tables for row in table.rows)
        return "\n".join(blocks)
    if suffix == ".pdf":
        reader = PdfReader(str(path), strict=False)
        return "\n".join((page.extract_text() or "") for page in reader.pages[:35])
    if suffix == ".xlsx":
        book = load_workbook(path, read_only=True, data_only=True); blocks = []
        for sheet in book.worksheets[:12]:
            for row in sheet.iter_rows(max_row=250, max_col=30, values_only=True):
                values = [str(value) for value in row if value not in (None, "")]
                if values: blocks.append(" | ".join(values))
        return "\n".join(blocks)
    return ""


def clean(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\x00", " ")).strip()


def paragraphs(text: str) -> list[str]:
    parts = re.split(r"(?:\r?\n)+|(?<=[.!?])\s+(?=[A-ZÁÉÍÓÚÑ])", text)
    return [clean(item) for item in parts if 120 <= len(clean(item)) <= 1_100]


def privacy_blocked(text: str) -> bool:
    return bool(EMAIL.search(text) or PHONE.search(text) or PERSON_ID.search(text) or IBAN.search(text) or PERSON_CUES.search(text))


def overlap(a: str, b: str) -> float:
    left = set(re.findall(r"\w+", a.lower())); right = set(re.findall(r"\w+", b.lower()))
    return len(left & right) / max(1, len(left | right))


def select_candidates(records: list[dict]) -> dict[str, dict | None]:
    chosen = {}
    for key, (_, terms) in FIELDS.items():
        ranked = []
        for record in records:
            low = record["text"].lower(); hits = sum(term in low for term in terms)
            if not hits or privacy_blocked(record["text"]) or PLACEHOLDER.search(record["text"]): continue
            score = hits * 8 + min(len(record["text"]), 800) / 100 + (2 if record["extension"] == ".docx" else 0)
            ranked.append((score, record))
        ranked.sort(key=lambda item: item[0], reverse=True)
        selected = []
        for score, record in ranked:
            if any(overlap(record["text"], item["text"]) > .72 for item in selected): continue
            selected.append({**record, "score": round(score, 2)})
            if len(selected) == 3: break
        chosen[key] = selected[0] if selected else None
    return chosen


def institutional_facts(records: list[dict], entity_name: str) -> dict:
    token = entity_name.split()[0].lower()
    tax = defaultdict(list); address_groups = defaultdict(lambda: {"records": [], "postcodes": Counter(), "street": "", "number": ""})
    for record in records:
        text = record["text"]
        if token not in text.lower(): continue
        for match in re.finditer(r"\b[A-Z]\d{8,9}\b", text):
            window = text[max(0, match.start()-250):match.end()+250]
            if token in window.lower(): tax[match.group().upper()].append(record)
        if re.search(r"direcci[oó]n social|domicilio|sede", text, re.I):
            for match in re.finditer(r"(?:Calle|C/|Avda\.?|Avenida)\s+([^,;|]{2,70}?),?\s+(\d{1,4})(?:\s+[^;|]{0,40}?\b(\d{5})\b)?", text, re.I):
                if abs(match.start() - text.lower().find(token)) > 300: continue
                street = clean(match.group(1)).strip("., "); number = match.group(2); postcode = match.group(3)
                key = (re.sub(r"\W+", "", street.lower()), number)
                address_groups[key]["records"].append(record); address_groups[key]["street"] = street; address_groups[key]["number"] = number
                if postcode: address_groups[key]["postcodes"][postcode] += 1

    addresses = {}
    for group in address_groups.values():
        postcode = group["postcodes"].most_common(1)[0][0] if group["postcodes"] else ""
        value = f"Calle {group['street']}, {group['number']}" + (f", C.P. {postcode}" if postcode else "")
        addresses[value] = group["records"]

    def proposal(values: dict, label: str):
        if not values: return {"field": label, "status": "missing", "value": None, "conflicts": [], "evidence": []}
        ordered = sorted(values.items(), key=lambda item: len(item[1]), reverse=True); value, evidence = ordered[0]
        return {"field": label, "status": "proposed" if len(evidence) >= 2 else "review_required", "value": value,
                "conflicts": [item[0] for item in ordered[1:]], "evidence": [item["source"] for item in evidence[:5]]}

    return {
        "legal_name": {"field": "Razón social", "status": "proposed", "value": entity_name, "conflicts": [], "evidence": []},
        "tax_id": proposal(tax, "Identificador fiscal"),
        "registered_address": proposal(addresses, "Domicilio social"),
    }


def main():
    config = args(); corpus = config.corpus.resolve(); out = config.out.resolve(); out.mkdir(parents=True, exist_ok=True)
    records = []; counts = Counter(); supported = {".docx", ".pdf", ".xlsx"}
    for path in corpus.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in supported: continue
        try: text = extract(path)
        except Exception: counts["extraction_errors"] += 1; continue
        digest = hashlib.sha256(path.read_bytes()).hexdigest()[:12]; relative = str(path.relative_to(corpus))
        for paragraph in paragraphs(text):
            records.append({"text": paragraph, "source": relative, "document_id": digest, "extension": path.suffix.lower()})
        counts[path.suffix.lower()] += 1

    selected = select_candidates(records); facts = institutional_facts(records, config.entity_name)
    filled = sum(value is not None for value in selected.values()); metrics = {
        "documents_read": sum(counts.values()) - counts["extraction_errors"], "paragraphs_considered": len(records),
        "filled_sections": filled, "total_sections": len(FIELDS), "coverage_percent": round(100 * filled / len(FIELDS)),
        "external_ai_calls": 0, "personal_paragraphs_excluded": sum(privacy_blocked(item["text"]) for item in records)
    }
    report = {"tenant": config.tenant, "created_at": datetime.now(timezone.utc).isoformat(), "privacy": {
        "processing": "local_only", "external_ai_calls": 0, "cross_tenant_reuse": False,
        "human_review_required": True, "personal_data_policy": "excluded_by_default"
    }, "metrics": metrics, "institutional_facts": facts, "section_proposals": selected}
    report_path = out / "master-facts-proposals.json"; docx_path = out / "plantilla-maestra-propuesta.docx"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    build_docx(docx_path, config.entity_name, facts, selected, metrics, FIELDS)
    print(json.dumps({"ok": True, "metrics": metrics, "report": str(report_path), "docx": str(docx_path)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
